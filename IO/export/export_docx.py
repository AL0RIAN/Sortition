__all__ = ["ExportPairsDOCX", "Participant", "Pair"]

from docx import Document  # Main python DOCX class
from docx.shared import Pt, Cm  # Points, Centimeters
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # Paragraph alignment
from docx.oxml.shared import OxmlElement, qn  # Magic
from os import path as os_path  # Working with external files

'''
Sample data.
Contains test data, that can be exported in DOCX file.

:param start-date - The day tournament is starting
:param end-date - The day tournament is ending
:param pairs - A pair of two participants
    :param is-circle-trinity - If there's three participants in circle trinity, then True, else False
    :param age-category - Info about pair participants age category
    :param weight-category - Info about pair participants weight category
    :param red/blue - Contestant
        :param last-name - Last name of contestant
        :param first-name - First name of contestant
        :param club - The club contestant is in


For sample purpose only.
'''
sample_data = {
    "start-date": "21.04.2022",
    "end-date": "23.03.2022",
    "pairs": [
        {
            "is-circle-trinity": True,
            "age-category": "8-9",
            "weight-category": "44",
            "red": {
                "last-name": "Виенко",
                "first-name": "Богдан",
                "club": "Саньда"
            },

            "blue": {
                "last-name": "Рублев",
                "first-name": "Дмитрий",
                "club": "Патриот"
            }
        },

        {
            "is-circle-trinity": False,
            "age-category": "8-9",
            "weight-category": "44",
            "red": {
                "last-name": "Виенко",
                "first-name": "Богдан",
                "club": "Саньда"
            },

            "blue": {
                "last-name": "Рублев",
                "first-name": "Дмитрий",
                "club": "Патриот"
            }
        }
    ]
}


class Participant:
    """
    Class that represents info about particular contestant

    :param last_name: Last name of contestant
    :param first_name: First name of contestant
    :param club: The club contestant is in
    """

    def __init__(self, last_name: str, first_name: str, club: str):
        """

        :param last_name: Last name of contestant
        :param first_name: First name of contestant
        :param club: The club contestant is in
        """
        self.last_name = last_name
        self.first_name = first_name
        self.club = club


class Pair:
    """
    Class that represents info about particular pair of contestants

    :param is-circle-trinity: If there's three participants in circle trinity, then True, else False
    :param age-category: Info about pair participants age category
    :param weight-category: Info about pair participants weight category
    :param red/blue: Contestant
    """

    def __init__(self, is_circle_trinity: bool, age_category: str, weight_category: str,
                 red: Participant, blue: Participant):
        """

        :param is_circle_trinity: If there's three participants in circle trinity, then True, else False
        :param age_category: Info about pair participants age category
        :param weight_category: Info about pair participants weight category
        :param red: Contestant
        :param blue: Contestant
        """
        self.is_circle_trinity = is_circle_trinity
        self.age_category = age_category
        self.weight_category = weight_category
        self.red = red
        self.blue = blue


class ExportPairsDOCX:
    """Main DOCX export class, that converts parsed info into DOCX couples file."""

    __DEFAULT_DATA = {
        "start-date": "",
        "end-date": "",
        "pairs": []
    }

    def __init__(self, start_date: str = "{{start-date}}", end_date: str = "{{end-date}}",
                 data: dict = __DEFAULT_DATA):
        """

        :param start_date: The day tournament is starting
        :param end_date: The day tournament is ending
        :param data: Contains data, that will be exported in DOCX file.
        """
        if __name__ == "__main__":  # If this file is run, then base.docx is in the same folder as this file
            self.__document = Document(os_path.abspath("base.docx"))
        else:
            # Else base.docx is in the same folder as this file,
            # but because this file is imported, you have to do this
            self.__document = Document(os_path.abspath("export\\base.docx"))
        self.start_date = start_date
        self.end_date = end_date
        self.__table = self.__document.tables[0]  # Only one table is needed
        self.__data = data

    def __print_dates(self) -> None:
        """
        Prints dates in the start of exporting file

        :return: None
        """
        for paragraph in self.__document.paragraphs:
            for run in paragraph.runs:
                if run.text == "{{start-date}}":
                    run.text = self.start_date

                if run.text == "{{end-date}}":
                    run.text = self.end_date

    def __print_table(self) -> None:
        """
        Prints containing in __data info into file table

        :return: None
        """
        count = 1  # Left table column

        for pair in self.__data["pairs"]:
            row = self.__table.add_row()
            row.height = Cm(0.6)
            row_cells = row.cells

            row_cells[0].text = str(count)
            row_cells[0].paragraphs[0].runs[0].font.name = "Century Gothic"
            row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
            row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            row_cells[1].text = f"{pair['age-category']} лет\n{pair['weight-category']} кг"
            row_cells[1].paragraphs[0].runs[0].font.bold = True
            row_cells[1].paragraphs[0].runs[0].font.size = Pt(9)
            row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            '''RED'''
            row_cells[2].text = f"{pair['red']['last-name']} {pair['red']['first-name']}"
            row_cells[2].paragraphs[0].runs[0].font.size = Pt(10)
            row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            row_cells[3].text = pair['red']['club']
            row_cells[3].paragraphs[0].runs[0].font.size = Pt(9)
            row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            row_cells[4].text = "Кругов." if pair['is-circle-trinity'] else ""
            row_cells[4].paragraphs[0].runs[0].font.bold = True
            row_cells[4].paragraphs[0].runs[0].font.size = Pt(8)
            row_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            '''BLUE'''
            row_cells[5].text = f"{pair['blue']['last-name']} {pair['blue']['first-name']}"
            row_cells[5].paragraphs[0].runs[0].font.size = Pt(10)
            row_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            row_cells[6].text = pair['blue']['club']
            row_cells[6].paragraphs[0].runs[0].font.size = Pt(9)
            row_cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            for i in range(7):
                row_cells[i].paragraphs[0].paragraph_format.space_after = Cm(0)

                '''
                Кирилл или Ярик, если вы это смотрите, я сам не знаю как работает это колесо сансары
                '''
                tc = row_cells[i]._tc
                tcPr = tc.get_or_add_tcPr()
                tcValign = OxmlElement('w:vAlign')
                tcValign.set(qn('w:val'), "center")
                tcPr.append(tcValign)

            count += 1

    def add_pair(self, pair: Pair) -> None:
        """
        Adds one pair to __data.

        :param pair: instance of Pair class.
        :return: None
        """

        new_pair = {
            "is-circle-trinity": False,
            "age-category": "",
            "weight-category": "",
            "red": {
                "last-name": "",
                "first-name": "",
                "club": ""
            },

            "blue": {
                "last-name": "Рублев",
                "first-name": "Дмитрий",
                "club": "Патриот"
            }
        }

        new_pair["is-circle-trinity"] = pair.is_circle_trinity
        new_pair["age-category"] = pair.age_category
        new_pair["weight-category"] = pair.weight_category

        new_pair["red"]["last-name"] = pair.red.last_name
        new_pair["red"]["first-name"] = pair.red.first_name
        new_pair["red"]["club"] = pair.red.club

        new_pair["blue"]["last-name"] = pair.blue.last_name
        new_pair["blue"]["first-name"] = pair.blue.first_name
        new_pair["blue"]["club"] = pair.blue.club

        self.__data["pairs"].append(new_pair)

    def export(self, path="") -> None:
        """
        Congratulations! You're exporting a file. Leave path empty to export in root folder

        :param path: Path where to save export file
        :return: None
        """
        self.__print_dates()
        self.__print_table()
        if path != "":
            self.__document.save(os_path.abspath(f"{path}/Состав пар турнир {self.start_date} - {self.end_date}.docx"))
        else:
            self.__document.save(f"Состав пар турнир {self.start_date} - {self.end_date}.docx")

class ExportRequestDOCX:
    """DOCX export class, that converts dictionary info from open window into DOCX participants file."""

    __AGE_CATEGORY_PATTERN = {  # Pattern cell-value: converted-value
        "10-": "Юнаки молодшого віку до 10 років",
        "10-11": "Юнаки молодшого віку 10-11 років",
        "12-13": "Юнаки старшого віку 12-13 років",
        "14-15": "Юніори 14-15 років",
        "16-17": "Юніори 16-17 років",
        "18+": "Дорослі 18+ років"
    }

    def __init__(self, info_dict: dict, region_name: str = "{{REGION-NAME}}") -> None:

        self.__info_dict = info_dict
        self.__document = Document(os_path.abspath("request_base.docx"))
        self.__table = self.__document.tables[0]
        self.__participant_num_counter = 1
        self.region_name = region_name


    def __add_header(self, text: str):

        row = self.__table.add_row()
        row.cells[0].merge(row.cells[-1])

        row.cells[0].text = ExportRequestDOCX.__AGE_CATEGORY_PATTERN[text]
        row.cells[0].paragraphs[0].runs[0].font.name = "Times New Roman"
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].italic = True
        row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def __add_participant(self, participant_info: list):

        row = self.__table.add_row()
        row.height = Cm(0.45)
        row_cells = row.cells

        row_cells[0].text = str(self.__participant_num_counter)
        row_cells[0].paragraphs[0].runs[0].font.name = "Times New Roman"
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        row_cells[1].text = participant_info[0]
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(10)

        for i in range(2, 13):
            row_cells[i].text = participant_info[i-1]
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for i in range(13):
            row_cells[i].paragraphs[0].paragraph_format.space_after = Cm(0)

            '''
            Кирилл или Ярик, если вы это смотрите, я все еще не знаю как работает это колесо сансары
            '''
            tc = row_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            tcValign = OxmlElement('w:vAlign')
            tcValign.set(qn('w:val'), "center")
            tcPr.append(tcValign)

        self.__participant_num_counter += 1

    def __analyze(self):
        for i in range(len(self.__info_dict.keys())):
            if len(self.__info_dict[i]) == 1:
                self.__add_header(self.__info_dict[i][0])

            else:
                self.__add_participant(self.__info_dict[i])

    def __print_region(self):

        for paragraph in self.__document.paragraphs:
            for run in paragraph.runs:
                if run.text == "{{REGION-NAME":
                    run.text = self.region_name

                if run.text == "}}":
                    run.text = ""

    def export(self, path="", file_name="noname"):

        self.__print_region()
        self.__analyze()

        if path != "":
            self.__document.save(os_path.abspath(f"{path}/{file_name}.docx"))
        else:
            self.__document.save(f"{file_name}.docx")

data = {0: ['18+'], 1: ['Логвиненко Дмитро', '13.05.03', 'М', 'КМС', '63.0', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Запорожець В.Г.'], 2: ['16-17'], 3: ['Вершина Дмитро', '22.10.04', 'М', 'І', '62.2', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Марченко В.А.'], 4: ['Кравченко Данило', '08.01.05', 'М', 'КМС', '82.0', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Запорожець В.Г.'], 5: ['14-15'], 6: ['Карнаух Дмитро ', '06.01.06', 'М', 'КМС', '50.0', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Запорожець В.Г.'], 7: ['Цимбал Андрій', '28.12.05', 'М', 'КМС', '54.3', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Марченко В.А.'], 8: ['Малик Іван', '25.12.06', 'М', 'Іюн', '58.2', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Запорожець В.Г.'], 9: ['Троценко Даніїл', '28.02.07', 'М', 'Іюн', '47.0', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Запорожець В.Г.'], 10: ['12-13'], 11: ['Кадиров Кирило', '09.01.09', 'М', 'Іюн', '34.4', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Марченко В.А.'], 12: ['Лиско Ігор', '20.04.09', 'М', 'Іюн', '43.5', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Марченко В.А.'], 13: ['Клочкова Софія', '26.05.09', 'Ж', 'б/р', '56.7', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Марченко В.А.'], 14: ['Колєсник Владислав', '23.04.08', 'М', 'Іюн', '49.6', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Запорожець В.Г.'], 15: ['Рябчиков Тимофій', '08.04.08', 'М', 'б/р', '38.2', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Запорожець В.Г.'], 16: ['Ткаченко Микита', '27.07.08', 'М', 'ІІІ', '62.2', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Марченко В.А.'], 17: ['10-11'], 18: ['Веретюк Іванна ', '07.08.11', 'Ж', 'ІІюн', '29.0', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Марченко В.А.'], 19: ['Павленко Богдан', '18.09.11', 'М', 'ІІІюн', '30.0', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Колдовський П.'], 20: ['Чередніченко Денис', '15.01.11', 'М', 'ІІІюн', '30.7', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Марченко В.А.'], 21: ['Сафронова Вікторія', '22.02.11', 'Ж', 'ІІюн', '31.5', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Марченко В.А.'], 22: ['Фєчікін Руслан', '17.08.10', 'М', 'ІІІюн', '39.5', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Марченко В.А.'], 23: ['Талалаєв Дмитро', '20.10.11', 'М', 'б/р', '37.8', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Колдовський П.'], 24: ['Федюков Артур', '23.12.09', 'М', 'Іюн', '39.3', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Запорожець В.Г.'], 25: ['Лізогубов Данило', '17.08.11', 'М', 'ІІюн', '51.5', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Марченко В.А.'], 26: ['Луговський Ян', '14.03.10', 'М', 'ІІюн', '52.1', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Запорожець В.Г.'], 27: ['10-'], 28: ['Борін Артем', '06.09.12', 'М', 'ІІІюн', '36.0', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Колдовський П.'], 29: ['Баранов Нікіта', '23.01.13', 'М', 'ІІІюн', '29.4', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Запорожець В.Г.'], 30: ['Бобров Мирон', '05.09.12', 'М', 'ІІІюн', '30.9', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Запорожець В.Г.'], 31: ['Лисенко Денис', '23.06.12', 'М', 'ІІІюн', '25.0', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Запорожець В.Г.'], 32: ['Ляшенко Єгор', '06.08.12', 'М', 'б/р', '31.0', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Запорожець В.Г.'], 33: ['Алавідзе Михайло', '05.06.12', 'М', 'ІІІюн', '40.2', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Запорожець В.Г.'], 34: ['Бруй Марк', '24.04.12', 'М', 'ІІІюн', '32.3', '+', '', '', '', 'Запоріжжя', 'ДЮСШ 10', 'Марченко В.А.']}
ExportRequestDOCX(data, region_name="Запорізької області").export(file_name="Test")

# TODO: ExportRefereeProtocolDOCX
