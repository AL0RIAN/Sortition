from docx import Document


class Parser:
    """Main parser of DOCX files"""

    # Number of cells in a table row
    __DOC_TABLE_WIDTH = 13

    # Pattern table-cell: cell value
    __TABLE_PATTERN = {
        1: "name",
        2: "birthday",
        3: "gender",
        4: "category",
        5: "weight",
        6: "is_sanda",
        7: "is_cinda",
        8: "is_tuishou",
        9: "is_vinchun",
        10: "region",
        11: "club",
        12: "trainer",

    }

    # Pattern cell-value: converted-value
    __AGE_CATEGORY_PATTERN = {
        "Юнаки молодшого віку до 10 років": "10-",
        "Юнаки молодшого віку 10-11 років": "10-11",
        "Юнаки старшого віку 12-13 років": "12-13",
        "Юніори 14-15 років": "14-15",
        "Юніори 16-17 років": "16-17",
        "Дорослі 18+ років": "18+"
    }

    def __init__(self, file_name: str):
        """

        :param file_name: Import file name
        """

        self.__data = dict(competition_region="", participants=[])
        self.file_name = file_name
        self.__doc = Document(file_name)
        self.__table = self.__doc.tables[0]
        self.__data["competition_region"] = self.__parse_competition_region()
        self.__parse_participants_info()

    def __parse_competition_region(self) -> str:
        """
        Collects competition region info

        :return: str
        """

        for paragraph in self.__doc.paragraphs:
            if "Команди: " in paragraph.text:
                return paragraph.text.replace("Команди: ", "")

    def __is_age_category(self, cell) -> bool:
        """
        If collected data is age category infom then True, else False

        :param cell: cell from table
        :return: bool
        """
        return True if cell.text in Parser.__AGE_CATEGORY_PATTERN else False

    def __insert_info_in_dict(self, dict: dict, row: int, col: int) -> None:
        """
        Searches the cell data type, and inserts particular value to particular key in dict

        :param dict: participant info
        :param row: row in table
        :param col: column in table
        :return: None
        """
        cell = self.__table.rows[row].cells[col]
        dict[Parser.__TABLE_PATTERN[col]] = cell.text

    def __parse_participants_info(self) -> None:
        """
        Main parse loop

        :return: None
        """

        # Because 1 in row is number of participant
        row = 2

        while True:
            try:
                participant_info = {}

                cell = self.__table.rows[row].cells[1]

                if self.__is_age_category(cell=cell):
                    self.__data["participants"].insert(0, {"age_category": Parser.__AGE_CATEGORY_PATTERN[cell.text],
                                                           "data": []})
                    row += 1
                    continue

                participant_info["name"] = cell.text

                for i in range(1, Parser.__DOC_TABLE_WIDTH):
                    self.__insert_info_in_dict(participant_info, row, i)

                self.__data["participants"][0]["data"].append(participant_info)

                row += 1

            except IndexError:
                break

    def result(self) -> dict:
        """
        Returns the result of parsing in dictionary format

        :return: dict
        """
        # return json.dumps(self.__data, ensure_ascii=False)
        return self.__data
